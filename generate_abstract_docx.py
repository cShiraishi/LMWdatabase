import docx
import sys

doc = docx.Document()
doc.add_heading("Mapeamento Quimioinformático da Megabiodiversidade Brasileira: AmazoniaDatabase e Blue_AmazonDB em Prol da Ciência, do Ecoturismo e do Bem-Estar no Prato", 0)

# Intro
p1 = doc.add_paragraph()
p1.add_run("Introdução: ").bold = True
p1.add_run("A intersecção entre gastronomia e ecoturismo demanda sistemas fundamentados no modelo holístico One Health. Embora o Brasil detenha megabiodiversidade terrestre e costeira, a conversão deste patrimônio botânico e marinho em sistemas alimentares resilientes e atrativos turísticos requer validação científico-tecnológica estruturada.")

# Obj
p2 = doc.add_paragraph()
p2.add_run("Objetivo(s): ").bold = True
p2.add_run("Descrever e analisar as bases de dados quimioinformáticas AmazoniaDatabase e Blue_AmazonDB (integradas à plataforma Life's Molecule Warehouse), desenvolvidas para o mapeamento estrutural e a catalogação de metabólitos da biodiversidade brasileira com aplicação em gastronomia funcional.")

# Metodologia
p3 = doc.add_paragraph()
p3.add_run("Metodologia: ").bold = True
p3.add_run("Investigou-se e indexou-se in silico compostos derivados de matrizes biológicas. Foram calculados parâmetros físico-químicos (massa molecular, LogP) e preditores de absorção (Regra dos Cinco de Lipinski - Ro5). As bioatividades foram extraídas da literatura paritária e estruturadas em interfaces visuais para a correlação cruzada de dados moleculares com propriedades botânicas/marinhas nutricionais.")

# Resultados
p4 = doc.add_paragraph()
p4.add_run("Resultados: ").bold = True
p4.add_run("Foram catalogados 560 registros moleculares na plataforma global. A AmazoniaDatabase compilou 360 bioativos provenientes de 18 frutos endêmicos, com predominância representativa em Euterpe oleracea (61 compostos polares, sobretudo antocianinas antioxidantes), Rollinia mucosa (alcaloides lipofílicos de atividade citotóxica) e Myrciaria dubia. A conformidade com a Ro5 no acervo terrestre foi de 36,4%, valor tecnicamente esperado dada a conformação molecular glicosídica e polifenólica na qual doadores e aceitadores de hidrogênio (HBD/HBA) excedem os limiares apolares convencionais. Na demografia bioprospetiva, quantificou-se ampla intersecção molecular entre os potenciais antioxidante (218 menções) e anti-inflamatório (215). Paralelamente, a Blue_AmazonDB validou estruturalmente 200 compostos marinhos ('Blue Foods', incluindo diterpenos e haliclonacilaminas), apresentando um perfil farmacológico com conformidade Ro5 significativamente superior (66,0%) e LogP médio de 3,20.")

# Conclusoes
p5 = doc.add_paragraph()
p5.add_run("Conclusões: ").bold = True
p5.add_run("A compilação sistematizada in silico consubstancia evidência quimioinformática robusta sobre o arcabouço funcional de ingredientes nativos e potenciais novos bioprodutos. A integração quantitativa destes parâmetros farmacológicos em plataformas de acesso aberto suporta o desenvolvimento lógico de diretrizes culinárias sustentáveis, validando a transição da biodiversidade brasileira para o turismo gastronômico com estrito rigor perante a 'Saúde Única'.")

doc.save("abstract_gastronomia_eshte.docx")
